#!/usr/bin/env python3
"""
Generate the Chillibyte document templates as real .docx files.

Run:  build/venv/bin/python build/make_docx.py
      (or point at any venv with python-docx installed)

Templates 1-7 come from the existing Google Docs template pack.
Templates 8-9 are new, closing gaps G4 (sales handover) and G2 (variation order).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), '..', 'templates', 'docx')

# Brand, from the live theme
RED   = RGBColor(0xED, 0x1B, 0x24)
GREEN = RGBColor(0x8D, 0xC5, 0x3E)
DARK  = RGBColor(0x21, 0x23, 0x33)
BODY  = RGBColor(0x40, 0x40, 0x42)
MUTED = RGBColor(0x6A, 0x6A, 0x6A)

FONT = 'Aptos'          # sensible modern default; Articulat CF is not installable here
FONT_FALLBACK = 'Calibri'


# --------------------------------------------------------------------------- helpers

def shade(cell, hexcolour):
    tc = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolour)
    tc.append(el)


def new_doc(title, subtitle, gate=None, route=None, is_new=False):
    doc = Document()

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = BODY
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.paragraph_format.space_after = Pt(6)

    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # Brand bar
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(2)
    r = bar.add_run('CHILLIBYTE')
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RED
    r2 = bar.add_run('  ·  WEBSITE DELIVERY PROCESS')
    r2.font.size = Pt(8)
    r2.font.bold = True
    r2.font.color.rgb = MUTED

    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    hr = h.add_run(title)
    hr.font.size = Pt(22)
    hr.font.bold = True
    hr.font.color.rgb = DARK

    # The gap-review additions are part of the process now, so the templates
    # no longer flag themselves as new. The parameter is kept so the call
    # sites still read correctly and provenance stays traceable in the script.
    _ = is_new

    sub = doc.add_paragraph()
    sr = sub.add_run(subtitle)
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED
    sr.italic = True

    meta = []
    if gate:
        meta.append(gate)
    if route:
        meta.append(f'Routes: {route}')
    meta.append('Version: v0.1 draft → v1.0 on approval → v1.1 after a logged change')
    m = doc.add_paragraph()
    mr = m.add_run('   |   '.join(meta))
    mr.font.size = Pt(8)
    mr.font.bold = True
    mr.font.color.rgb = RED
    m.paragraph_format.space_after = Pt(14)

    return doc


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK
    return p


def guidance(doc, text):
    """Italic instruction line. Tell the user to delete it — templates that keep
    their own instructions in the final document look unfinished to a client."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = MUTED
    return p


def fields(doc, items, width_label=Cm(5.4)):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    for label, hint in items:
        row = t.add_row()
        c0, c1 = row.cells
        c0.width = width_label
        p = c0.paragraphs[0]
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK
        shade(c0, 'F5F5F5')
        if hint:
            p1 = c1.paragraphs[0]
            r1 = p1.add_run(hint)
            r1.font.size = Pt(8.5)
            r1.italic = True
            r1.font.color.rgb = MUTED
    doc.add_paragraph()
    return t


def table(doc, columns, rows=3, hints=None):
    t = doc.add_table(rows=1, cols=len(columns))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, c in enumerate(columns):
        p = hdr[i].paragraphs[0]
        r = p.add_run(c)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], '212333')
    if hints:
        cells = t.add_row().cells
        for i, hint in enumerate(hints):
            p = cells[i].paragraphs[0]
            r = p.add_run(hint)
            r.font.size = Pt(8)
            r.italic = True
            r.font.color.rgb = MUTED
        rows -= 1
    for _ in range(max(0, rows)):
        t.add_row()
    doc.add_paragraph()
    return t


def bullets(doc, items, style='List Bullet'):
    for i in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(i)
        r.font.size = Pt(10)


def callout(doc, title, text, colour=GREEN):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    shade(c, 'EDF4E3' if colour == GREEN else 'FBE7E8')
    p = c.paragraphs[0]
    r = p.add_run(title + '  ')
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = DARK
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = BODY
    doc.add_paragraph()
    return t


def save(doc, name):
    path = os.path.abspath(os.path.join(OUT, name))
    doc.save(path)
    print(f'  {name}')


# --------------------------------------------------------------------------- 1. Sales handover

def sales_handover():
    d = new_doc('Sales handover note',
                'Completed by sales, accepted by delivery, before discovery is scheduled.',
                gate='Gate 0', route='light / standard / full', is_new=True)

    callout(d, 'Why this exists.',
            'What sales promised is the commonest single origin of scope disagreement. Twenty minutes here '
            'protects every stage that follows. The exclusions list is the most valuable part of this document.')

    h2(d, 'Commercial basis')
    fields(d, [
        ('Client', ''), ('Project name', ''), ('Sold by', ''), ('Handover date', ''),
        ('Contract value', ''), ('Payment structure', 'e.g. 40% on order, 30% at design approval, 30% on launch'),
        ('Target launch', 'State whether this is contractual or aspirational'),
        ('Route sold as', 'light / standard / full'),
    ])

    h2(d, 'What was sold — inclusions')
    guidance(d, 'List explicitly. Take these from the proposal, not from memory.')
    bullets(d, ['', '', '', '', ''])

    h2(d, 'What was NOT sold — exclusions')
    guidance(d, 'This is the important half. "Not included: copywriting, photography, translation, '
                'ongoing SEO" is worth more than any other line in this document. A blank exclusions '
                'list is the single strongest predictor of a scope argument later.')
    bullets(d, ['', '', '', '', '', ''])

    h2(d, 'Revision allowance agreed at sale')
    fields(d, [
        ('Design concepts included', 'A number, not "a couple"'),
        ('Revision rounds included', 'On the chosen concept'),
        ('Staging feedback rounds', 'Usually one consolidated round'),
    ])

    h2(d, 'Verbal commitments made during the sale')
    guidance(d, 'Anything said in a meeting that the client may reasonably have taken as a promise. '
                'Capture it here and either honour it or formally retract it now.')
    table(d, ['Commitment', 'Made by', 'Honour or retract?'], rows=3)

    h2(d, 'Known commercial risks passed to delivery')
    table(d, ['Risk', 'Why it matters', 'Owner'], rows=3)

    h2(d, 'Delivery acceptance')
    fields(d, [
        ('Accepted by', 'Head of Web'), ('Date', ''),
        ('Gap between sold and deliverable',
         'If what was sold cannot be delivered for the price, say so NOW. It gets harder every week.'),
        ('Resolution', 'Repriced / rescoped / accepted as-is with the shortfall recorded'),
        ('Basecamp reference', 'Gate 0: sales handover accepted'),
    ])
    save(d, 'sales-handover-note.docx')


# --------------------------------------------------------------------------- 2. Discovery summary

def discovery_summary():
    d = new_doc('Discovery summary',
                'The foundation artefact. Every later document traces back to an answer in here.',
                gate='Gate 1', route='light / standard / full')

    h2(d, 'Project')
    fields(d, [
        ('Client', ''), ('Project name', ''), ('Discovery date', ''), ('Discovery lead', ''),
        ('Account manager', ''), ('Client stakeholder representative', ''),
        ('Project route', 'light / standard / full — record the reason in the decision log as DEC-001'),
    ])

    h2(d, 'Business context')
    fields(d, [
        ('What the organisation does', ''),
        ('Why the website is being created or changed', ''),
        ('Current website position', ''),
        ('Commercial or operational context', ''),
    ])

    h2(d, 'Primary goals')
    guidance(d, 'State goals so they can be judged met or not met. "Look more modern" is not a goal; '
                '"generate enquiries from specifying engineers rather than general procurement" is.')
    table(d, ['Goal', 'How we would know it worked'], rows=3)

    h2(d, 'Target audiences')
    table(d, ['Audience', 'Needs', 'Concerns', 'Likely conversion route'], rows=3)

    h2(d, 'Core services, products or offers')
    table(d, ['Service / offer', 'Description', 'Priority', 'Required page or section'], rows=3)

    h2(d, 'Stakeholder concerns')
    table(d, ['Concern', 'Raised by', 'Impact on structure, content, design or build'], rows=3)

    h2(d, 'Content ownership')
    guidance(d, 'Assign an owner per section, not "the client". Vague content ownership is the most '
                'reliable predictor of a late project.')
    table(d, ['Content', 'Owner (named person)', 'Status', 'Deadline'], rows=4)

    h2(d, 'Brand and design inputs')
    fields(d, [
        ('Existing brand assets', ''), ('Competitors or references', ''),
        ('Visual preferences', ''),
        ('Visual dislikes', 'Ask directly. The "avoid" list saves more time than the inspiration list.'),
        ('Photography or image assets', ''),
    ])

    h2(d, 'Technical context')
    fields(d, [
        ('Existing WordPress site', ''), ('Hosting', ''),
        ('Domains and DNS', 'Who actually controls the registrar? It is rarely who you expect.'),
        ('Forms', 'Where do submissions go, and who reads them?'),
        ('Integrations', ''), ('Analytics / search tools', ''),
        ('Multilingual requirements', ''), ('Known plugin requirements', ''),
        ('Known technical risks', ''),
    ])

    h2(d, 'Guardrails — agreed at Gate 1')
    callout(d, 'These take minutes now and are close to impossible to impose later.',
            'Do not skip this section on small projects. It is where the process pays for itself.')
    fields(d, [
        ('Named client approver', 'One person. Others may comment; only this person approves.'),
        ('Approval response window', 'e.g. 5 working days. Silence past it pauses the project — it does not mean yes.'),
        ('Browser and device support matrix', 'Which browsers, devices and breakpoints. This is what makes QA finishable.'),
        ('Accessibility target', 'e.g. WCAG 2.2 AA. "Accessible" is not a testable statement.'),
        ('Revision allowance', 'Carried forward from the sales handover'),
        ('Estimate baseline', 'From the spec builder. Core hours + contingency. The yardstick for all later change.'),
    ])

    h2(d, 'Assumptions')
    table(d, ['Assumption', 'Impact if wrong', 'Owner', 'Validation required'], rows=3)

    h2(d, 'Risks')
    table(d, ['Risk', 'Likelihood', 'Impact', 'Mitigation', 'Owner'], rows=4)

    h2(d, 'Exclusions')
    guidance(d, 'Carry forward from the sales handover and add anything discovery has ruled out. '
                'The spec builder generates a client-ready version of this list.')
    table(d, ['Item', 'Reason excluded', 'Future phase candidate?'], rows=4)

    h2(d, 'Immediate next steps')
    table(d, ['Action', 'Owner', 'Due date', 'Basecamp reference'], rows=3)
    save(d, 'discovery-summary.docx')


# --------------------------------------------------------------------------- 3. Page outline

def page_outline():
    d = new_doc('Page outline',
                'What a page must say and do, before any layout decision is made.',
                gate='Gate 4', route='light / standard / full')

    callout(d, 'Outline templates, not every page.',
            'A repeatable template needs one outline, not twenty. Outline the unique pages and each '
            'distinct template.')

    h2(d, 'Page')
    fields(d, [
        ('Page name', ''), ('Proposed URL', ''), ('Parent page', ''),
        ('Template type', 'unique / repeatable / standard content page'),
        ('Primary audience', ''), ('Business purpose', 'What the business needs this page to achieve'),
        ('User purpose', 'What the visitor came here to do'),
        ('Primary CTA', ''), ('Secondary CTA', ''),
        ('Content owner', 'Named person'), ('Content deadline', 'A date. Not "ASAP".'),
        ('SEO notes', ''),
    ])

    h2(d, 'Section structure')
    guidance(d, 'Write section purposes as jobs — "answer the objection about lead times", not "text block". '
                'Approximate lengths prevent both a one-line section and an unplanned essay, and they tell '
                'the client exactly how much writing they have committed to.')
    table(d,
          ['Section', 'Purpose', 'Approx. length', 'Media required', 'Likely block', 'Owner'],
          rows=8,
          hints=['H1 + intro', 'Name the thing and the problem it solves',
                 '40–60 words', 'Product image', 'Hero', 'Client'])

    h2(d, 'FAQs and objections')
    guidance(d, 'Source these from what the sales team is actually asked on the phone.')
    table(d, ['Question', 'Answer intent', 'Source / owner'], rows=4)

    h2(d, 'Open decisions')
    guidance(d, 'List them rather than resolving them by quietly picking one option.')
    table(d, ['Decision needed', 'Options', 'Owner', 'Decision log ID'], rows=3)
    save(d, 'page-outline.docx')


# --------------------------------------------------------------------------- 4. Block specification

def block_spec():
    d = new_doc('Block specification',
                'One per block being adapted or built new. Reused blocks do not need one.',
                gate='Gate 5', route='standard / full')

    h2(d, 'Block')
    fields(d, [
        ('Block name', ''),
        ('Block type', 'existing / existing with variation / new reusable / one-off'),
        ('Used on', 'List every page. If it is only one page, ask whether it should be a block at all.'),
        ('Purpose', ''), ('Content role', ''), ('Design role', ''), ('Build priority', ''),
    ])

    h2(d, 'Reuse decision')
    fields(d, [
        ('Existing block available?', ''),
        ('Reuse / adapt / build new', ''),
        ('Reason', 'Record in the decision log — this is a decision with a reversal cost'),
        ('Build impact', 'low / medium / high, and an hours figure'),
    ])

    h2(d, 'Fields')
    table(d, ['Field name', 'Field type', 'Required', 'Default', 'Notes'], rows=6,
          hints=['Heading', 'Text', 'Yes', '—', 'Max 60 chars'])

    h2(d, 'Display and responsive rules')
    guidance(d, 'Responsive behaviour is a structural decision. Agree it here, not at QA.')
    fields(d, [
        ('Desktop layout', ''), ('Tablet layout', ''), ('Mobile layout', ''),
        ('Image behaviour', ''), ('Text limits', ''), ('Button / link behaviour', ''),
        ('Animation / motion', ''),
        ('Accessibility notes', 'Focus states, heading level, keyboard reachability, contrast'),
    ])

    h2(d, 'Structure and styling separation')
    callout(d, 'The reuse principle.',
            'Block PHP structure and field patterns are shared and reusable. The visual system varies '
            'through theme SCSS. Every fork is a permanent maintenance cost and a lost reuse opportunity.')
    fields(d, [
        ('PHP structure', ''), ('field group', ''),
        ('Base SCSS', 'Shared, structural'), ('Theme SCSS', 'Project-specific appearance'),
        ('Design library entry required', 'yes / no'),
    ])

    h2(d, 'Design notes')
    fields(d, [
        ('Fixed structure', 'What the designer must not change'),
        ('Open visual questions', 'Where the designer\'s judgement is wanted'),
        ('Required states or variants', ''),
        ('References', ''), ('Avoid', ''),
    ])

    h2(d, 'Build notes')
    fields(d, [('Dependencies', ''), ('Known risks', ''), ('Testing notes', '')])
    save(d, 'block-specification.docx')


# --------------------------------------------------------------------------- 5. Designer brief

def designer_brief():
    d = new_doc('Designer brief and style direction',
                'Approved structure and principles. What is fixed, and where judgement is wanted.',
                gate='Gate 7', route='standard / full')

    callout(d, 'The point of this document.',
            'Separating what is fixed from what is open is not a constraint on the designer — it is a '
            'licence. It says exactly where their judgement is wanted, and stops effort going into an '
            'area that is already approved.')

    h2(d, 'Project')
    fields(d, [('Project', ''), ('Client', ''), ('Prepared by', ''), ('Date', ''),
               ('Design route', 'light / standard / full'),
               ('Concepts and revisions included', 'From the sales handover')])

    h2(d, 'Context')
    fields(d, [('Business summary', ''), ('Website purpose', ''), ('Audience summary', ''),
               ('Primary conversion goals', '')])

    h2(d, 'Approved artefacts this design must work from')
    guidance(d, 'Name the version. The designer is working from these specific versions, and a later '
                'change to any of them is a change log entry.')
    table(d, ['Artefact', 'Version', 'Approved at'], rows=6,
          hints=['Sitemap', 'v1.0', 'Gate 3'])

    h2(d, 'What is FIXED — approved, do not change')
    bullets(d, ['Sitemap and URL structure', 'Page section order per the outlines',
                'The block model and block inventory', 'Content hierarchy and heading levels',
                'Technical constraints', ''])

    h2(d, 'What is OPEN — your judgement is wanted')
    bullets(d, ['Visual language and overall expression', 'Colour refinement within the agreed roles',
                'Typographic scale and hierarchy', 'Image treatment and cropping',
                'Spacing, rhythm and density', 'Component polish, states and interaction detail', ''])

    callout(d, 'The control.',
            'Improve visual hierarchy, rhythm, expression and polish freely. Do not change sitemap, '
            'content strategy, block scope or technical structure without a logged decision. Good ideas '
            'are welcome; silent ones are not. If a change is worth making, log it and price it.', RED)

    h2(d, 'Design principles')
    guidance(d, 'Three to five, specific to this client. A principle you could paste into any other '
                'project is not a principle.')
    bullets(d, ['', '', ''])

    h2(d, 'Colour direction')
    guidance(d, 'Define by role, not by hex. Roles survive "can we try it in blue"; values do not.')
    fields(d, [('Primary action', ''), ('Emphasis', ''), ('Recessive surface', ''),
               ('Semantic / status', ''),
               ('Contrast expectations', 'State the ratio. Body text needs 4.5:1 against its background.'),
               ('Avoid', '')])

    h2(d, 'Typography direction')
    fields(d, [('Display / headline role', ''), ('Body role', ''), ('UI / table role', ''),
               ('Readability requirements', 'Minimum sizes, tabular figures if there are data tables'),
               ('Avoid', '')])

    h2(d, 'Image and media direction')
    fields(d, [('Preferred image style', ''), ('Treatment', ''),
               ('Icon / illustration approach', ''), ('Avoid', '')])

    h2(d, 'Per-block design tasks')
    table(d, ['Block', 'Design task', 'Constraints', 'Variants needed'], rows=6,
          hints=['Spec table', 'Real design thought needed', 'Must read at 375px', 'Compact + full'])

    h2(d, 'Review criteria — agreed before work starts')
    guidance(d, 'Agreeing these up front is what makes the internal review objective. '
                '"Does it respect the block model" has an answer; "do we like it" has a mood.')
    bullets(d, [
        'Does the design follow the approved structure?',
        'Does it respect the block model, or has new scope been created?',
        'Does it improve hierarchy without adding unapproved components?',
        'Is the visual system coherent across the site?',
        'Are reused blocks treated consistently wherever they appear?',
        'Does every text/background pair meet the accessibility target?',
        'Can it be built efficiently in the agreed WordPress stack?',
    ])
    save(d, 'designer-brief.docx')


# --------------------------------------------------------------------------- 6. Build specification

def build_spec():
    d = new_doc('Build specification',
                'Approved design translated into templates, blocks, fields and a re-baselined estimate.',
                gate='Gate 10', route='standard / full')

    h2(d, 'Project')
    fields(d, [('Project', ''), ('Prepared by', ''), ('Date', ''),
               ('Approved design reference', 'Name and version'),
               ('Approved sitemap reference', 'Name and version'),
               ('Approved block inventory reference', 'Name and version')])

    h2(d, 'Stack and environments')
    guidance(d, 'Record this properly. It is what allows a second developer to join without a briefing — '
                'which is the whole parallel-working argument.')
    fields(d, [('WordPress stack', ''), ('Theme / child theme', ''), ('Block framework version', ''),
               ('Plugins approved', 'No others without a logged decision'),
               ('GitHub repository', ''),
               ('Branching approach', 'e.g. main → staging → feature branches; no direct commits to main'),
               ('Environments', 'Local, dev, staging, production')])

    h2(d, 'Page and template list')
    table(d, ['Page', 'URL', 'Template', 'Blocks used', 'Content owner', 'Notes'], rows=6)

    h2(d, 'Block build list')
    table(d, ['Block', 'Action', 'fields', 'PHP file', 'SCSS file', 'Library entry', 'Est. hrs'],
          rows=6, hints=['Hero', 'Adapt', 'heading, image, cta', 'blocks/hero/', 'theme/_hero.scss', 'Yes', '5'])

    h2(d, 'field groups')
    table(d, ['Field group', 'Location rule', 'Fields', 'Notes'], rows=4)

    h2(d, 'Content entry plan')
    guidance(d, 'Content entry is real effort and is routinely left out of estimates. '
                'Populating thirty pages is a cost — estimate it explicitly.')
    table(d, ['Page group', 'Pages', 'Content status', 'Image status', 'Owner', 'Est. hrs'], rows=5)

    h2(d, 'Estimate re-baseline')
    callout(d, 'Keep the variance.',
            'The difference between the Gate 1 baseline and this re-baseline is one of the most useful '
            'numbers the process produces. Project after project it shows where estimating is '
            'systematically wrong — and that is worth more than any single estimate.')
    table(d, ['', 'Gate 1 baseline', 'Gate 10 re-baseline', 'Variance'], rows=5,
          hints=['Core process steps', '167 hrs', '171 hrs', '+4'])

    h2(d, 'Access and credentials')
    guidance(d, 'Verify each one works. "They sent a password" is not access, and a dead credential '
                'found in launch week is a launch delay.')
    table(d, ['System', 'Needed for', 'Client owner', 'Requested', 'Received', 'Verified'], rows=6)

    h2(d, 'Technical risks')
    table(d, ['Risk', 'Impact', 'Mitigation', 'Owner'], rows=4)

    h2(d, 'QA focus')
    fields(d, [('Support matrix', 'From Gate 1 — QA tests against this and nothing beyond it'),
               ('Accessibility target', 'From Gate 1'),
               ('Forms', ''), ('Navigation', ''), ('Performance', ''),
               ('SEO basics', ''), ('Analytics', ''),
               ('Redirects', 'Test against the real old URL list from the audit')])
    save(d, 'build-specification.docx')


# --------------------------------------------------------------------------- 7. Gate ticket

def gate_ticket():
    d = new_doc('Basecamp gate ticket',
                'Paste into Basecamp. One ticket per gate. This is what approval actually looks like.',
                gate='Every gate', route='light / standard / full')

    callout(d, 'Approval is not "looks good".',
            'It means the named artefacts, at the named versions, are accepted as the basis for the next '
            'stage. Without a named version there is nothing to point at when it is disputed later.')

    h2(d, 'Ticket')
    fields(d, [
        ('Gate', 'Number and name'),
        ('Purpose', 'What this gate confirms'),
        ('Artefacts reviewed', 'Name AND version for each one'),
        ('Decision', 'Approved / approved with notes / not approved / deferred'),
        ('Notes', 'Conditions, caveats, follow-up points'),
        ('Impact of approval', 'What work is now allowed to proceed'),
        ('If not approved', 'Where the work returns to, and what needs re-approval'),
        ('Change control note',
         'Any future change to these approved artefacts goes in the change log and is assessed for impact.'),
        ('Decision log references', 'DEC-000'),
        ('Change log references', 'CHG-000'),
        ('Variation references', 'VAR-000'),
        ('Approved by', 'Name and role'),
        ('Date', ''),
    ])

    h2(d, 'Ready-to-paste template')
    guidance(d, 'Copy the block below into the Basecamp ticket body and fill it in.')
    p = d.add_paragraph()
    r = p.add_run(
        'GATE [n]: [name]\n\n'
        'Artefacts being approved:\n'
        '- [Artefact name] [version]\n'
        '- [Artefact name] [version]\n\n'
        'Approving this confirms these artefacts are the agreed basis for the next stage.\n\n'
        'This unlocks: [what work can now start]\n\n'
        'Any later change to these items can still be considered, but will be logged and '
        'reviewed for impact on structure, content, design, build and timing.\n\n'
        'Approved by: ____________________   Date: __________'
    )
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = DARK
    save(d, 'gate-ticket.docx')


# --------------------------------------------------------------------------- 8. Client approval note

def approval_note():
    d = new_doc('Client approval note',
                'The client-facing version. Sanitised, short, and specific about what is being agreed.',
                gate='Client gates', route='light / standard / full')

    h2(d, 'Approval')
    p = d.add_paragraph()
    r = p.add_run('The following are approved as the basis for the next stage of work:')
    r.font.size = Pt(10.5)
    bullets(d, ['[Artefact name and version]', '[Artefact name and version]', '[Artefact name and version]'])

    p = d.add_paragraph()
    r = p.add_run(
        'This approval confirms the agreed direction is suitable for the next project phase. '
        'Any later change to these approved items can still be considered, but will be reviewed for '
        'its impact on structure, content, design, build and timing — and where it adds work, we will '
        'agree that with you before proceeding.')
    r.font.size = Pt(10.5)

    d.add_paragraph()
    fields(d, [('Approved by', ''), ('Role', ''), ('Date', ''), ('Basecamp reference', '')])

    h2(d, 'What happens next')
    fields(d, [('Now starting', ''), ('We need from you', 'What, from whom, by when'),
               ('Next approval point', 'Which gate, and roughly when')])

    callout(d, 'Note on tone.',
            'The final sentence of the approval paragraph is doing commercial work. It is phrased as '
            'collaboration rather than as a warning, but it establishes that additions are discussed '
            'and agreed rather than assumed.')
    save(d, 'client-approval-note.docx')


# --------------------------------------------------------------------------- 9. Variation order

def variation_order():
    d = new_doc('Variation order',
                'Turns an accepted change into agreed, billable work — before the work happens.',
                gate='Raised from the change log', route='light / standard / full', is_new=True)

    callout(d, 'Why this exists.',
            'The change log recorded what changed and whether it was accepted, but carried no effort '
            'figure and no route to a charge. A change process that cannot raise a charge simply '
            'converts scope creep into unpaid work more politely.', RED)

    h2(d, 'Variation')
    fields(d, [
        ('Variation reference', 'VAR-000'),
        ('Date raised', ''), ('Raised by', ''),
        ('Related change log entry', 'CHG-000'),
        ('Related approved artefact', 'What this changes, and at which version it was approved'),
        ('Gate affected', ''),
    ])

    h2(d, 'What is being requested')
    fields(d, [('Description', ''), ('Reason given', ''),
               ('Classification', 'Correction (free) / preference (within allowance) / new scope (this form)')])

    guidance(d, 'Only new scope reaches this document. Corrections are fixed at no charge; preferences '
                'are absorbed within the revision allowance. Classifying honestly is what makes the '
                'conversation reasonable rather than defensive.')

    h2(d, 'Impact assessment')
    table(d, ['Area', 'Impact', 'Hours'], rows=7,
          hints=['Structure / sitemap', 'New page and navigation entry', '3'])

    fields(d, [('Total additional effort', 'Hours — not "some impact"'),
               ('Rate', ''), ('Total price', ''),
               ('Schedule impact', 'Days added to the timeline, and which gate moves'),
               ('Effect on the estimate baseline', 'Accepted variations update the baseline')])

    h2(d, 'Options offered to the client')
    guidance(d, 'Offering a choice is what makes this a conversation rather than a bill. All three of '
                'these are legitimate outcomes.')
    bullets(d, [
        'Proceed now at the price above — the timeline moves by the stated amount',
        'Defer to after launch as a separate piece of work',
        'Do not proceed — the approved scope stands unchanged',
    ])

    h2(d, 'Decision')
    fields(d, [('Client decision', 'Proceed / defer / decline'),
               ('Approved by', 'Named client approver'), ('Date', ''),
               ('Basecamp reference', ''),
               ('If absorbed at no charge',
                'Record it anyway. Absorbing twenty minutes as goodwill is a choice; absorbing twenty '
                'hours because nobody was counting is not.')])
    save(d, 'variation-order.docx')


# --------------------------------------------------------------------------- run

if __name__ == '__main__':
    os.makedirs(OUT, exist_ok=True)
    print('Generating document templates:')
    sales_handover()
    discovery_summary()
    page_outline()
    block_spec()
    designer_brief()
    build_spec()
    gate_ticket()
    approval_note()
    variation_order()
    print('Done.')
